import streamlit as st
import os
import io
import tempfile
from dotenv import load_dotenv
from groq import Groq
import PyPDF2
import docx
from PIL import Image
import pytesseract
try:
    from pdf2image import convert_from_path
except ImportError:
    st.warning("pdf2image is not installed. OCR functionality might be limited.")
    # Fallback function to avoid errors if pdf2image is not installed
    def convert_from_path(*args, **kwargs):
        return []

# Load environment variables
load_dotenv()

# Initialize Groq client
client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

def extract_text_from_pdf(file):
    """Extracts text from a PDF file."""
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        page_text = page.extract_text()
        text += page_text
    
    # Check if text extraction failed or returned very little text
    if len(text.strip()) < 100:  # Adjust threshold as needed
        st.warning("Standard text extraction yielded minimal results. Attempting OCR...")
        text = extract_text_using_ocr(file)
    
    return text

def extract_text_using_ocr(file):
    """Extracts text from images/scanned PDFs using OCR."""
    text = ""
    with tempfile.TemporaryDirectory() as path:
        pdf_path = os.path.join(path, "temp.pdf")
        
        # Save the uploaded file
        with open(pdf_path, "wb") as f:
            file.seek(0)
            f.write(file.read())
        file.seek(0)  # Reset file pointer for future operations
        
        try:
            images = convert_from_path(pdf_path)
            for image in images:
                text += pytesseract.image_to_string(image)
        except Exception as e:
            st.error(f"OCR error: {str(e)}")
    
    return text

def extract_text_from_docx(file):
    """Extracts text from a Word document."""
    doc = docx.Document(file)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    extracted_text = '\n'.join(full_text)
    
    # Check if extracted text is minimal
    if len(extracted_text.strip()) < 100:  # Adjust threshold as needed
        st.warning("Standard text extraction yielded minimal results from DOCX. Attempting OCR on document images...")
        extracted_text = extract_text_from_docx_images(file) or extracted_text
    
    return extracted_text

def extract_text_from_docx_images(file):
    """Extract text from images embedded in a DOCX file using OCR."""
    try:
        # Save uploaded file to temporary directory
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_file_path = os.path.join(temp_dir, "temp.docx")
            
            # Save the uploaded file
            with open(temp_file_path, "wb") as f:
                file.seek(0)
                f.write(file.read())
            file.seek(0)  # Reset file pointer
            
            # Load the document again
            doc = docx.Document(temp_file_path)
            
            # Extract and process images
            extracted_text = ""
            
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # Get image data
                        image_data = rel.target_part.blob
                        
                        # Create a PIL Image from binary data
                        image = Image.open(io.BytesIO(image_data))
                        
                        # Use OCR to extract text
                        image_text = pytesseract.image_to_string(image)
                        if image_text.strip():
                            extracted_text += image_text + "\n\n"
                    except Exception as e:
                        st.error(f"Error processing image in DOCX: {str(e)}")
                        continue
            
            return extracted_text
    except Exception as e:
        st.error(f"Error extracting images from DOCX: {str(e)}")
        return ""

def extract_text_from_txt(file):
    """Extracts text from a text file."""
    text = file.read().decode('utf-8')
    return text

def extract_text(file, file_type):
    """Extract text from different file types."""
    if file_type == "pdf":
        return extract_text_from_pdf(file)
    elif file_type == "docx":
        return extract_text_from_docx(file)
    elif file_type == "txt":
        return extract_text_from_txt(file)
    else:
        return "Unsupported file type"

def extract_resume_data(resume_text):
    """Extracts structured data from resume text using Groq API."""
    prompt = f"""Extract ONLY the following information from the resume text provided below:
    - Full Name
    - Email Address
    - Phone Number
    - Location (City, State/Country)
    - Education (Degree, Institution, Year) - list all
    - Work Experience (Company, Position, Duration) - list all
    - Skills (Technical and Soft Skills) - list all (only names like python,nextjs,leadership,etc)
    - Years of Experience - IMPORTANT: If not explicitly stated, calculate this by adding up all work experience durations or estimate based on career progression just show the number no explaination needed

    Resume Text:
    {resume_text}

    Return ONLY the extracted information in this exact format - do not include any additional information, analysis, or commentary:

    ## Full Name
    [Extracted name]

    ## Email Address
    [Extracted email]

    ## Phone Number
    [Extracted phone]

    ## Location
    [Extracted location]

    ## Education
    - [Degree], [Institution], [Year]
    - [Additional education entries]

    ## Work Experience
    - [Company], [Position], [Duration]
    - [Additional work experience entries]

    ## Skills
    [List of extracted skills]

    ## Years of Experience
    [Number of years] - YOU MUST PROVIDE THIS! Calculate if not explicitly stated in resume aand display only the number

    If a field is not found, indicate "Not found" for that field only.
    """

    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        ],
        model="llama3-70b-8192", # Using LLaMA 3 70B model
        temperature=0.2, # Lower temperature for more consistent and precise output
        max_tokens=1000 # Limit response length to avoid unnecessary content
    )
    return chat_completion.choices[0].message.content

# Streamlit App
st.title("Resume Processing System")

uploaded_file = st.file_uploader("Upload a resume", type=["pdf", "docx", "txt"])

if uploaded_file is not None:
    st.write("File Uploaded Successfully!")
    
    # Get file type
    file_type = uploaded_file.name.split('.')[-1].lower()
    
    # Extract text based on file type
    with st.spinner(f"Extracting text from {file_type.upper()} file..."):
        resume_text = extract_text(uploaded_file, file_type)
    st.subheader("Extracted Text:")
    st.text_area("Resume Content", resume_text, height=300)

    # Process with Groq
    if st.button("Extract Information"):
        if not os.environ.get("GROQ_API_KEY") or os.environ.get("GROQ_API_KEY") == "YOUR_GROQ_API_KEY":
            st.error("GROQ_API_KEY not configured. Please set it in the .env file.")
        elif not resume_text.strip():
            st.warning("The extracted text is empty. Cannot process.")
        else:
            with st.spinner("Extracting information using Groq LLaMA-3..."):
                extracted_data = extract_resume_data(resume_text)
            st.subheader("Extracted Information:")
            st.markdown(extracted_data)
else:
    st.info("Please upload a PDF, DOCX, or TXT file to begin.")

