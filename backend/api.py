import os
import io
import os
import tempfile
import json
import uuid
import logging
import asyncio
import concurrent.futures
import re
from datetime import datetime
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Depends, Query, Request
from fastapi.responses import JSONResponse, RedirectResponse
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv
from groq import Groq
import PyPDF2
import docx
from PIL import Image
import pytesseract
import uvicorn
from typing import List, Optional, Dict, Any
from pydantic import BaseModel
from enum import Enum

# Define enums early
class DuplicateHandling(str, Enum):
    STRICT = "strict"  # Block both file and content duplicates
    ALLOW_UPDATES = "allow_updates"  # Allow content duplicates (updated resumes)
    ALLOW_ALL = "allow_all"  # Allow all uploads (no duplicate checking)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import database module
from database import (
    get_db, Candidate, Education, Skill, WorkExperience, Status, init_db, 
    save_candidate_data, get_all_candidates, shortlist_candidate,
    calculate_file_hash, check_duplicate_file, check_duplicate_candidate_content, 
    generate_batch_id, save_candidate_data_with_hash
)
from sqlalchemy.orm import Session

# Import S3 storage module
from s3_storage import upload_file_to_s3, generate_presigned_url

# Import shortlisting service
from shortlisting_service import shortlist_candidates, CandidateScore, ShortlistingResult

# Import chat service
from chat_service import chat_service, ChatRequest, ChatResponse, ChatHistory

# Import authentication middleware
from auth_middleware import get_current_user, get_current_user_optional

try:
    from pdf2image import convert_from_path
except ImportError:
    print("pdf2image is not installed. OCR functionality might be limited.")
    # Fallback function to avoid errors if pdf2image is not installed
    def convert_from_path(*args, **kwargs):
        return []

# Load environment variables
load_dotenv()

# Initialize Groq client
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
client = Groq(api_key=GROQ_API_KEY)

app = FastAPI()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, replace with your frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def extract_text_from_pdf(file_path):
    """Extracts text from a PDF file."""
    pdf_reader = PyPDF2.PdfReader(file_path)
    text = ""
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        page_text = page.extract_text()
        text += page_text
    
    # Check if text extraction failed or returned very little text
    if len(text.strip()) < 100:  # Adjust threshold as needed
        print("Standard text extraction yielded minimal results. Attempting OCR...")
        text = extract_text_using_ocr(file_path)
    
    return text

def extract_text_using_ocr(file_path):
    """Extracts text from images/scanned PDFs using OCR."""
    text = ""
    try:
        images = convert_from_path(file_path)
        for image in images:
            text += pytesseract.image_to_string(image)
    except Exception as e:
        print(f"OCR error: {str(e)}")
    
    return text

def extract_text_from_docx(file_path):
    """Extracts text from a Word document."""
    doc = docx.Document(file_path)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    extracted_text = '\n'.join(full_text)
    
    # Check if extracted text is minimal
    if len(extracted_text.strip()) < 100:  # Adjust threshold as needed
        print("Standard text extraction yielded minimal results from DOCX. Attempting OCR on document images...")
        extracted_text = extract_text_from_docx_images(file_path) or extracted_text
    
    return extracted_text

def extract_text_from_docx_images(file_path):
    """Extract text from images embedded in a DOCX file using OCR."""
    try:
        # Load the document
        doc = docx.Document(file_path)
        
        # Extract and process images
        extracted_text = ""
        
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    # Get image data
                    image_data = rel.target_part.blob
                    
                    # Create a PIL Image from binary data
                    image = Image.open(io.BytesIO(image_data))
                    
                    # Use OCR to extract text
                    image_text = pytesseract.image_to_string(image)
                    if image_text.strip():
                        extracted_text += image_text + "\n\n"
                except Exception as e:
                    print(f"Error processing image in DOCX: {str(e)}")
                    continue
        
        return extracted_text
    except Exception as e:
        print(f"Error extracting images from DOCX: {str(e)}")
        return ""

def extract_text_from_txt(file_path):
    """Extracts text from a text file."""
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
    return text

def extract_resume_data(resume_text):
    prompt = f"""Extract ONLY the following information from the resume text provided below:
    - Full Name
    - Email Address
    - Phone Number
    - Location (City, State/Country)
    - Education (Degree, Institution, Year) - list all
    - Work Experience (Company, Position, Duration) - list all
    - Skills (Technical and Soft Skills) - list all individual skills separated by commas (e.g., "Python, JavaScript, Leadership")
    - Years of Experience - IMPORTANT: If not explicitly stated, calculate this by adding up all work experience durations or estimate based on career progression just show the number no explaination needed

    Resume Text:
    {resume_text}

    Return ONLY the extracted information in this exact format - do not include any additional information, analysis, or commentary:

    ## Full Name
    [Extracted name]

    ## Email Address
    [Extracted email]

    ## Phone Number
    [Extracted phone]

    ## Location
    [Extracted location]

    ## Education
    - [Degree], [Institution], [4-digit Year ONLY like 2020, 2019, etc]
    - [Additional education entries]

    ## Work Experience
    - [Company], [Position], [Duration]
    - [Additional work experience entries]

    ## Skills
    [List all skills separated by commas, e.g., "Python, JavaScript, Leadership, Project Management"]

    ## Years of Experience
    [Number of years] - YOU MUST PROVIDE THIS! Calculate if not explicitly stated in resume and display only the number

    IMPORTANT: 
    - For education year, use ONLY 4-digit years (like 2020, 2019, 2018).
    - For skills, make sure to list individual skills separated by commas. Each skill should be 1-3 words maximum.
    - If a field is not found, indicate "Not found" for that field only.
    """

    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        ],
        model="llama3-70b-8192", # Using LLaMA 3 70B model
        temperature=0.2, # Lower temperature for more consistent and precise output
        max_tokens=1000 # Limit response length to avoid unnecessary content
    )
    return chat_completion.choices[0].message.content

class ResponseModel(BaseModel):
    extracted_text: str
    parsed_data: Optional[str] = None
    candidate_id: Optional[int] = None

class ParsedResumeData(BaseModel):
    full_name: str
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    education: List[Dict[str, Any]] = []
    work_experience: List[Dict[str, Any]] = []
    skills: List[str] = []
    years_experience: Optional[int] = None

# Batch processing models
class FileProcessingResult(BaseModel):
    filename: str
    status: str  # 'success', 'error', 'duplicate'
    candidate_id: Optional[int] = None
    extracted_text: Optional[str] = None
    parsed_data: Optional[str] = None
    message: Optional[str] = None
    existing_candidate_id: Optional[int] = None

class BatchProcessingResponse(BaseModel):
    batch_id: str
    total_files: int
    successful: int
    failed: int
    duplicates: int
    results: List[FileProcessingResult]

def parse_markdown_data(markdown_data: str) -> ParsedResumeData:
    """
    Parse the markdown text returned by the LLM into a structured format
    """
    data = {
        "full_name": "Unknown",
        "email": None,
        "phone": None,
        "location": None,
        "education": [],
        "work_experience": [],
        "skills": [],
        "years_experience": 0
    }
    
    # Split by sections (starting with ##)
    sections = markdown_data.split("## ")
    for section in sections:
        if not section.strip():
            continue
        
        # Get section name and content
        lines = section.strip().split("\n")
        section_name = lines[0].strip().lower()
        content = "\n".join(lines[1:]).strip()
        
        if "full name" in section_name:
            data["full_name"] = content if content != "Not found" else "Unknown"
        elif "email" in section_name:
            data["email"] = content if content != "Not found" else None
        elif "phone" in section_name:
            data["phone"] = content if content != "Not found" else None
        elif "location" in section_name:
            data["location"] = content if content != "Not found" else None        
        elif "education" in section_name:
            # Process education entries
            for edu_entry in content.strip().split('\n'):
                if edu_entry.startswith('- '):
                    edu_entry = edu_entry[2:].strip()  # Remove list marker
                    if ',' in edu_entry:
                        parts = [p.strip() for p in edu_entry.split(',')]
                        
                        # Extract year from the parts - look for 4-digit numbers
                        year = ""
                        for part in parts:
                            # Look for 4-digit year (1900-2099)
                            import re
                            year_match = re.search(r'\b(19|20)\d{2}\b', part)
                            if year_match:
                                year = year_match.group()
                                break
                        
                        edu_dict = {
                            "degree": parts[0] if len(parts) > 0 else "",
                            "institution": parts[1] if len(parts) > 1 else "",
                            "year": year
                        }
                        data["education"].append(edu_dict)
        elif "work experience" in section_name:
            # Process work experience entries
            for exp_entry in content.strip().split('\n'):
                if exp_entry.startswith('- '):
                    exp_entry = exp_entry[2:].strip()  # Remove list marker
                    if ',' in exp_entry:
                        parts = [p.strip() for p in exp_entry.split(',')]
                        exp_dict = {
                            "company": parts[0] if len(parts) > 0 else "",
                            "position": parts[1] if len(parts) > 1 else "",
                            "duration": parts[2] if len(parts) > 2 else ""
                        }
                        data["work_experience"].append(exp_dict)
        elif "skills" in section_name:
            # Process skills
            skills_text = content.replace('Not found', '').strip()
            if skills_text:
                # Process skills with improved splitting logic
                # First handle bullet point list format if present
                if '\n-' in skills_text:
                    for skill_line in skills_text.split('\n'):
                        if skill_line.startswith('- '):
                            skill_text = skill_line[2:].strip()
                            # Further split by commas if present
                            if ',' in skill_text:
                                for sub_skill in skill_text.split(','):
                                    clean_skill = sub_skill.strip()
                                    if clean_skill:
                                        data["skills"].append(clean_skill)
                            else:
                                if skill_text:
                                    data["skills"].append(skill_text)
                # Handle comma-separated list
                elif ',' in skills_text:
                    skills = [s.strip() for s in skills_text.split(',')]
                    data["skills"].extend([s for s in skills if s])
                # Handle space-separated list (potential issue case)
                else:
                    # Try to break apart space-separated multi-word skills
                    # This handles the case where the LLM outputs skills without commas
                    import re
                    # Look for common skill separators or patterns
                    potential_skills = re.split(r'\s+(?=[A-Z]|\d)|(?<=\w)\s+(?=[A-Z])', skills_text)
                    
                    # If this produced too many small fragments or just one big chunk, 
                    # try a more aggressive splitting approach
                    if len(potential_skills) <= 1 or all(len(s.split()) < 2 for s in potential_skills):
                        # Split on spaces but limit each skill to a few words
                        words = skills_text.split()
                        current_skill = []
                        
                        for word in words:
                            current_skill.append(word)
                            
                            # When we reach 2 words, consider it a complete skill
                            if len(current_skill) == 2:
                                data["skills"].append(' '.join(current_skill))
                                current_skill = []
                        
                        # Add any remaining words as the last skill
                        if current_skill:
                            data["skills"].append(' '.join(current_skill))
                    else:
                        # Use the regex split result
                        data["skills"].extend([s.strip() for s in potential_skills if s.strip()])
        elif "years of experience" in section_name:
            # Extract years of experience
            years_text = content.replace('Not found', '0').strip()
            # Try to extract just the number
            import re
            match = re.search(r'(\d+)', years_text)
            if match:
                data["years_experience"] = int(match.group(1))
            else:
                try:
                    data["years_experience"] = int(years_text)
                except ValueError:
                    data["years_experience"] = 0
    
    # Post-process skills to ensure they're not too long
    processed_skills = []
    for skill in data["skills"]:
        # If a skill is unreasonably long, it's probably multiple skills
        if len(skill) > 30 or len(skill.split()) > 3:
            # Split long skills by common separators
            sub_skills = re.split(r',|\s+and\s+|\s+&\s+|\s*\|\s*|\s+/\s+', skill)
            for sub in sub_skills:
                # Further split if still too long
                if len(sub) > 30 or len(sub.split()) > 3:
                    # Split into individual words or pairs
                    words = sub.split()
                    for i in range(0, len(words), 2):
                        if i + 1 < len(words):
                            processed_skills.append(f"{words[i]} {words[i+1]}")
                        else:
                            processed_skills.append(words[i])
                else:
                    clean_sub = sub.strip()
                    if clean_sub:
                        processed_skills.append(clean_sub)
        else:
            processed_skills.append(skill)
    
    # Replace the original skills with the processed ones
    # Remove duplicates while preserving order
    seen = set()
    data["skills"] = [x for x in processed_skills if not (x in seen or seen.add(x))]
    
    return ParsedResumeData(**data)

class CandidateResponse(BaseModel):
    candidate_id: int
    full_name: str
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    years_experience: Optional[int] = None
    status: str = "pending"  # will store the enum value as string
    created_at: Optional[str] = None  # ISO format string
    resume_available: bool = False
    original_filename: Optional[str] = None
    skills: List[str] = []

    class Config:
        from_attributes = True  # This enables ORM model parsing

@app.post("/upload-resume/", response_model=ResponseModel)
async def upload_resume(
    request: Request,
    file: UploadFile = File(...), 
    parse: bool = Form(False), 
    save_to_db: bool = Form(False),
    duplicate_handling: DuplicateHandling = Form(DuplicateHandling.STRICT),
    db: Session = Depends(get_db),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Upload and process a resume file. 
    Set parse=true to extract structured data from the resume.
    Set save_to_db=true to save the parsed data to the database.
    Set duplicate_handling to control how duplicates are handled:
    - strict: Block both file and content duplicates
    - allow_updates: Allow content duplicates (updated resumes from same person)
    - allow_all: Allow all uploads (no duplicate checking)
    """
    # Check if file extension is supported
    file_extension = file.filename.split('.')[-1].lower()
    if file_extension not in ["pdf", "docx", "txt"]:
        raise HTTPException(status_code=400, detail="Unsupported file format. Please upload a PDF, DOCX, or TXT file.")
    
    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_file:
        temp_file_path = temp_file.name
        content = await file.read()
        temp_file.write(content)
    
    try:
        # Calculate file hash for duplicate checking when saving to DB
        if save_to_db and duplicate_handling != DuplicateHandling.ALLOW_ALL:
            file_hash = calculate_file_hash(temp_file_path)
            if file_hash:
                # Check for file-based duplicates (same file uploaded by same user)
                duplicate_info = check_duplicate_file(file_hash, current_user['id'])
                if duplicate_info and duplicate_handling == DuplicateHandling.STRICT:
                    os.unlink(temp_file_path)
                    raise HTTPException(
                        status_code=409, 
                        detail=f"Identical file already exists for candidate '{duplicate_info['candidate_name']}' (uploaded on {duplicate_info['upload_date'][:10]})"
                    )
                    # For non-strict handling, we'll continue and update the record
        
        # Extract text based on file type
        if file_extension == "pdf":
            extracted_text = extract_text_from_pdf(temp_file_path)
        elif file_extension == "docx":
            extracted_text = extract_text_from_docx(temp_file_path)
        elif file_extension == "txt":
            extracted_text = extract_text_from_txt(temp_file_path)
        
        # Parse the resume if requested
        parsed_data = None
        parsed_structured_data = None
        candidate_id = None
        if parse and extracted_text.strip():
            parsed_data = extract_resume_data(extracted_text)
            parsed_structured_data = parse_markdown_data(parsed_data)
            
            # Check for content-based duplicates (similar candidate data) when saving to DB
            if save_to_db and duplicate_handling == DuplicateHandling.STRICT:
                content_duplicate_info = check_duplicate_candidate_content(parsed_structured_data.dict(), current_user['id'])
                if content_duplicate_info and content_duplicate_info['is_likely_same_person']:
                    os.unlink(temp_file_path)
                    raise HTTPException(
                        status_code=409, 
                        detail=f"Similar candidate '{content_duplicate_info['candidate_name']}' already exists ({content_duplicate_info['similarity_percentage']:.0f}% match). This appears to be an updated resume of the same person."
                    )
            
            # Save to database if requested
            if save_to_db:
                # Generate a unique filename but keep it flat without extra folders
                original_filename = file.filename
                timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
                unique_id = str(uuid.uuid4())[:8]  # Use shorter UUID
                filename_base = os.path.splitext(original_filename)[0]
                filename_ext = os.path.splitext(original_filename)[1]
                
                # Create a flatter S3 key structure
                s3_key = f"resumes/{filename_base}_{timestamp}_{unique_id}{filename_ext}"
                
                # Upload to S3
                success, s3_url = upload_file_to_s3(temp_file_path, s3_key)
                
                if not success:
                    raise HTTPException(status_code=500, detail=f"Failed to upload file to S3: {s3_url}")
                
                # Generate a presigned URL for temporary access
                presigned_success, presigned_url = generate_presigned_url(s3_key, expiration=3600*24)  # 24 hours
                
                # Get file hash for storage
                file_hash = calculate_file_hash(temp_file_path) if 'file_hash' not in locals() else file_hash
                
                try:
                    # Save to database with S3 information and file hash
                    candidate_id = save_candidate_data_with_hash(
                        parsed_structured_data.dict(),
                        resume_file_path=s3_key,
                        resume_s3_url=presigned_url if presigned_success else s3_url,
                        original_filename=original_filename,
                        file_hash=file_hash,
                        batch_id=None,  # Single upload, no batch
                        user_id=current_user['id']
                    )
                    
                    if candidate_id is None:
                        raise HTTPException(status_code=500, detail="Failed to save candidate data to database.")
                except Exception as e:
                    logger.error(f"Database error: {str(e)}")
                    raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")
        
        # Clean up temporary file
        os.unlink(temp_file_path)
        
        return {
            "extracted_text": extracted_text,
            "parsed_data": parsed_data,
            "candidate_id": candidate_id
        }
    
    except HTTPException:
        # Clean up temporary file in case of HTTP error
        if os.path.exists(temp_file_path):
            os.unlink(temp_file_path)
        raise  # Re-raise HTTP exceptions
    except Exception as e:
        # Clean up temporary file in case of other errors
        if os.path.exists(temp_file_path):
            os.unlink(temp_file_path)
        logger.error(f"Error processing file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

async def process_single_file(file: UploadFile, batch_id: str, user_id: int, parse: bool = True, 
                             save_to_db: bool = True, duplicate_handling: DuplicateHandling = DuplicateHandling.STRICT) -> FileProcessingResult:
    """
    Process a single file in a batch operation
    
    Args:
        file: The uploaded file
        batch_id: The batch ID for this upload session
        user_id: ID of the user uploading the file
        parse: Whether to parse the resume
        save_to_db: Whether to save to database
        duplicate_handling: How to handle duplicates
        
    Returns:
        FileProcessingResult: Result of processing this file
    """
    filename = file.filename
    
    try:
        # Check if file extension is supported
        file_extension = filename.split('.')[-1].lower()
        if file_extension not in ["pdf", "docx", "txt"]:
            return FileProcessingResult(
                filename=filename,
                status="error",
                message="Unsupported file format. Please upload a PDF, DOCX, or TXT file."
            )
          
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_file:
            temp_file_path = temp_file.name
            content = await file.read()
            temp_file.write(content)
          
        # Calculate file hash for duplicate checking
        file_hash = calculate_file_hash(temp_file_path)
        if not file_hash:
            os.unlink(temp_file_path)
            return FileProcessingResult(
                filename=filename,
                status="error",
                message="Failed to calculate file hash"
            )
        
        # Check for file-based duplicates (same file uploaded by same user) only if not allowing all
        if duplicate_handling != DuplicateHandling.ALLOW_ALL:
            duplicate_info = check_duplicate_file(file_hash, user_id)
            if duplicate_info:
                # Only block duplicates in strict mode, otherwise we will update the record
                if duplicate_handling == DuplicateHandling.STRICT:
                    os.unlink(temp_file_path)
                    return FileProcessingResult(
                        filename=filename,
                        status="duplicate",
                        message=f"Identical file already exists for candidate '{duplicate_info['candidate_name']}' (uploaded on {duplicate_info['upload_date'][:10]})",
                        existing_candidate_id=duplicate_info['candidate_id']
                    )
                # Continue processing for non-strict modes, we'll update the record
        
        # Extract text based on file type
        if file_extension == "pdf":
            extracted_text = extract_text_from_pdf(temp_file_path)
        elif file_extension == "docx":
            extracted_text = extract_text_from_docx(temp_file_path)
        elif file_extension == "txt":
            extracted_text = extract_text_from_txt(temp_file_path)
        
        # Parse the resume if requested
        parsed_data = None
        parsed_structured_data = None
        candidate_id = None
        if parse and extracted_text.strip():
            parsed_data = extract_resume_data(extracted_text)
            parsed_structured_data = parse_markdown_data(parsed_data)
              
            # Check for content-based duplicates (similar candidate data) only in strict mode
            if duplicate_handling == DuplicateHandling.STRICT:
                content_duplicate_info = check_duplicate_candidate_content(parsed_structured_data.dict(), user_id)
                if content_duplicate_info and content_duplicate_info['is_likely_same_person']:
                    os.unlink(temp_file_path)
                    return FileProcessingResult(
                        filename=filename,
                        status="duplicate",
                        message=f"Similar candidate '{content_duplicate_info['candidate_name']}' already exists ({content_duplicate_info['similarity_percentage']:.0f}% match). This appears to be an updated resume of the same person.",
                        existing_candidate_id=content_duplicate_info['candidate_id']
                    )
            
            # Save to database if requested
            if save_to_db:
                # Generate a unique filename but keep it flat without extra folders
                timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
                unique_id = str(uuid.uuid4())[:8]  # Use shorter UUID
                filename_base = os.path.splitext(filename)[0]
                filename_ext = os.path.splitext(filename)[1]
                
                # Create a flatter S3 key structure
                s3_key = f"resumes/{filename_base}_{timestamp}_{unique_id}{filename_ext}"
                
                # Upload to S3
                success, s3_url = upload_file_to_s3(temp_file_path, s3_key)
                
                if not success:
                    os.unlink(temp_file_path)
                    return FileProcessingResult(
                        filename=filename,
                        status="error",
                        message=f"Failed to upload file to S3: {s3_url}"
                    )
                
                # Generate a presigned URL for temporary access
                presigned_success, presigned_url = generate_presigned_url(s3_key, expiration=3600*24)  # 24 hours
                
                try:
                    # Save to database with file hash and batch ID
                    candidate_id = save_candidate_data_with_hash(
                        parsed_structured_data.dict(),
                        resume_file_path=s3_key,
                        resume_s3_url=presigned_url if presigned_success else s3_url,
                        original_filename=filename,
                        file_hash=file_hash,
                        batch_id=batch_id,
                        user_id=user_id
                    )
                    
                    # If an existing record was updated (by email or file hash), this will return the ID
                    if candidate_id is None:
                        os.unlink(temp_file_path)
                        return FileProcessingResult(
                            filename=filename,
                            status="error",
                            message="Failed to save candidate data to database."
                        )
                        
                except Exception as e:
                    os.unlink(temp_file_path)
                    logger.error(f"Error saving candidate: {str(e)}")
                    # Check for specific error types we can handle better
                    error_message = str(e)
                    if "Duplicate entry" in error_message and "file_hash" in error_message:
                        return FileProcessingResult(
                            filename=filename,
                            status="duplicate",
                            message=f"This file has already been uploaded (duplicate file hash)."
                        )
                    elif "Duplicate entry" in error_message and "email" in error_message:
                        return FileProcessingResult(
                            filename=filename,
                            status="duplicate",
                            message=f"A candidate with this email already exists."
                        )
                    else:
                        return FileProcessingResult(
                            filename=filename,
                            status="error",
                            message=f"Database error: {error_message[:100]}..."  # Truncate very long error messages
                        )
        
        # Clean up temporary file
        os.unlink(temp_file_path)
        
        return FileProcessingResult(
            filename=filename,
            status="success",
            candidate_id=candidate_id,
            extracted_text=extracted_text,
            parsed_data=parsed_data,
            message="Successfully processed"
        )
    
    except Exception as e:
        # Clean up temporary file in case of error
        if 'temp_file_path' in locals() and os.path.exists(temp_file_path):
            os.unlink(temp_file_path)
        logger.error(f"Error processing file {filename}: {str(e)}")
        return FileProcessingResult(
            filename=filename,
            status="error",
            message=f"Error processing file: {str(e)}"
        )

@app.post("/upload-resumes-batch/", response_model=BatchProcessingResponse)
async def upload_resumes_batch(
    request: Request,
    files: List[UploadFile] = File(...),
    parse: bool = Form(True),
    save_to_db: bool = Form(True),
    duplicate_handling: DuplicateHandling = Form(DuplicateHandling.STRICT),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Upload and process multiple resume files in batch.
    Files are processed in parallel for better performance.
    Set duplicate_handling to control how duplicates are handled:
    - strict: Block both file and content duplicates
    - allow_updates: Allow content duplicates (updated resumes from same person)
    - allow_all: Allow all uploads (no duplicate checking)
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")
    
    if len(files) > 50:  # Limit batch size
        raise HTTPException(status_code=400, detail="Maximum 50 files allowed per batch")    # Generate a unique batch ID
    batch_id = generate_batch_id()
    
    try:
        # Process files in parallel using asyncio.gather
        tasks = [process_single_file(file, batch_id, current_user['id'], parse, save_to_db, duplicate_handling) for file in files]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Convert any exceptions to error results
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                processed_results.append(FileProcessingResult(
                    filename=files[i].filename,
                    status="error",
                    message=f"Unexpected error: {str(result)}"
                ))
            else:
                processed_results.append(result)
        
        # Calculate statistics
        total_files = len(processed_results)
        successful = len([r for r in processed_results if r.status == "success"])
        failed = len([r for r in processed_results if r.status == "error"])
        duplicates = len([r for r in processed_results if r.status == "duplicate"])
        
        return BatchProcessingResponse(
            batch_id=batch_id,
            total_files=total_files,
            successful=successful,
            failed=failed,
            duplicates=duplicates,
            results=processed_results
        )
    
    except Exception as e:
        logger.error(f"Error in batch processing: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error in batch processing: {str(e)}")

@app.post("/parse-text/")
async def parse_text(text: str = Form(...)):
    """
    Parse resume text and extract structured information.
    """
    if not text.strip():
        raise HTTPException(status_code=400, detail="Text is empty")
    
    try:
        parsed_data = extract_resume_data(text)
        return {"parsed_data": parsed_data}
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error parsing text: {str(e)}")

@app.get("/")
def read_root():
    return {"message": "Resume Processing API is running! Use /upload-resume/ endpoint to process resumes."}

@app.get("/candidates/", response_model=List[Dict[str, Any]])
async def get_candidates(
    request: Request,
    limit: int = Query(100, description="Maximum number of candidates to return"),
    status: Optional[str] = Query(None, description="Filter by status (pending, shortlisted, rejected)"),
    min_experience: Optional[int] = Query(None, description="Minimum years of experience"),
    max_experience: Optional[int] = Query(None, description="Maximum years of experience"),
    skills: Optional[str] = Query(None, description="Comma-separated list of skills to filter by"),
    location: Optional[str] = Query(None, description="Location to filter by (partial match)"),
    company: Optional[str] = Query(None, description="Company name to filter by (partial match)"),
    position: Optional[str] = Query(None, description="Position/title to filter by (partial match)"),
    education: Optional[str] = Query(None, description="Education/degree to filter by (partial match)"),
    db: Session = Depends(get_db),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """Get a list of all candidates with comprehensive filtering options (user-specific)"""
    try:
        status_enum = None
        if status:
            try:
                status_enum = Status[status.upper()]
            except KeyError:
                raise HTTPException(status_code=400, detail=f"Invalid status: {status}")
        
        # Parse skills from comma-separated string
        skills_list = None
        if skills:
            skills_list = [skill.strip() for skill in skills.split(',') if skill.strip()]
        
        # Get candidates with comprehensive filtering
        candidates = get_all_candidates(
            limit=limit, 
            status=status_enum, 
            user_id=current_user['id'],
            min_experience=min_experience,
            max_experience=max_experience,
            skills=skills_list,
            location=location,
            company=company,
            position=position,
            education=education
        )
        
        # Convert SQLAlchemy objects to dictionaries with comprehensive data
        result = []
        for candidate in candidates:
            skills = [skill.skill_name for skill in candidate.skills] if candidate.skills else []
            
            # Get education data
            education_data = []
            if candidate.education:
                for edu in candidate.education:
                    education_data.append({
                        "degree": edu.degree,
                        "institution": edu.institution,
                        "graduation_year": edu.graduation_year,
                        "gpa": edu.gpa
                    })
            
            # Get work experience data
            work_experience = []
            if candidate.work_experiences:
                for exp in candidate.work_experiences:
                    work_experience.append({
                        "company": exp.company,
                        "position": exp.position,
                        "start_date": exp.start_date,
                        "end_date": exp.end_date,
                        "duration": exp.duration,
                        "description": exp.description
                    })
            
            candidate_dict = {
                "candidate_id": candidate.candidate_id,
                "full_name": candidate.full_name,
                "email": candidate.email,
                "phone": candidate.phone,
                "location": candidate.location,
                "years_experience": candidate.years_experience,
                "status": candidate.status.value if candidate.status else "pending",
                "created_at": candidate.created_at.isoformat() if candidate.created_at else None,
                "resume_available": bool(candidate.resume_file_path),
                "original_filename": candidate.original_filename,
                "skills": skills,
                "education": education_data,
                "work_experience": work_experience
            }
            result.append(candidate_dict)
        
        return result
        
    except Exception as e:
        logger.error(f"Error getting candidates: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/candidates/{candidate_id}/shortlist")
def shortlist_candidate_endpoint(
    request: Request,
    candidate_id: int,
    db: Session = Depends(get_db),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Mark a candidate as shortlisted (user-specific)
    """
    try:
        # Get candidate and verify ownership
        candidate = db.query(Candidate).filter(
            Candidate.candidate_id == candidate_id,
            Candidate.user_id == current_user['id']
        ).first()
        
        if not candidate:
            raise HTTPException(status_code=404, detail=f"Candidate with ID {candidate_id} not found or not accessible")
        
        # Update status
        candidate.status = Status.SHORTLISTED
        candidate.updated_at = datetime.utcnow()
        db.commit()
        
        return {"message": f"Candidate with ID {candidate_id} has been shortlisted"}
    
    except Exception as e:
        db.rollback()
        logger.error(f"Error shortlisting candidate: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error shortlisting candidate: {str(e)}")

@app.patch("/candidates/{candidate_id}/status")
def update_candidate_status(
    request: Request,
    candidate_id: int,
    status: str = Form(...),
    db: Session = Depends(get_db),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Update a candidate's status (pending, shortlisted, rejected) (user-specific)
    """
    try:
        # Validate status
        valid_statuses = ['pending', 'shortlisted', 'rejected']
        if status.lower() not in valid_statuses:
            raise HTTPException(status_code=400, detail=f"Invalid status. Must be one of: {', '.join(valid_statuses)}")
        
        # Get candidate and verify ownership
        candidate = db.query(Candidate).filter(
            Candidate.candidate_id == candidate_id,
            Candidate.user_id == current_user['id']
        ).first()
        
        if not candidate:
            raise HTTPException(status_code=404, detail=f"Candidate with ID {candidate_id} not found or not accessible")
        
        # Update status
        status_enum = Status[status.upper()]
        candidate.status = status_enum
        candidate.updated_at = datetime.utcnow()
        db.commit()
        
        # Convert candidate to dict for response
        skills = [skill.skill_name for skill in candidate.skills] if candidate.skills else []
        candidate_dict = {
            "candidate_id": candidate.candidate_id,
            "full_name": candidate.full_name,
            "email": candidate.email,
            "phone": candidate.phone,
            "location": candidate.location,
            "years_experience": candidate.years_experience,
            "status": candidate.status.value,
            "created_at": candidate.created_at.isoformat() if candidate.created_at else None,
            "resume_available": bool(candidate.resume_file_path),
            "original_filename": candidate.original_filename,
            "skills": skills
        }
        
        return candidate_dict
        
    except Exception as e:
        db.rollback()
        logger.error(f"Error updating candidate status: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error updating candidate status: {str(e)}")

@app.post("/init-db")
def initialize_database():
    """
    Initialize the database - create tables if they don't exist
    """
    try:
        init_db()
        return {"message": "Database initialized successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Database initialization failed: {str(e)}")

@app.get("/resumes/{candidate_id}/view")
async def view_resume(
    request: Request,
    candidate_id: int,
    db: Session = Depends(get_db),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Get a URL to view a candidate's resume (user-specific)
    """
    try:
        # Get candidate from database and verify ownership
        candidate = db.query(Candidate).filter(
            Candidate.candidate_id == candidate_id,
            Candidate.user_id == current_user['id']
        ).first()
        
        if not candidate:
            raise HTTPException(status_code=404, detail=f"Candidate with ID {candidate_id} not found or not accessible")
        
        if not candidate.resume_file_path:
            raise HTTPException(status_code=404, detail="No resume file found for this candidate")
        
        # Check if we need to generate a new presigned URL (if existing one is old)
        # In a production system, you might want to check if the URL is expired        # Generate a new presigned URL with 24 hour expiration for inline viewing
        success, url = generate_presigned_url(candidate.resume_file_path, expiration=3600*24, inline=True)
        
        if not success:
            raise HTTPException(status_code=500, detail="Failed to generate URL for resume")
        
        # Get file extension for frontend handling
        file_extension = None
        if candidate.original_filename:
            file_extension = candidate.original_filename.split('.')[-1].lower()
          # Return the URL and filename information with enhanced metadata
        response_data = {
            "resume_url": url,
            "filename": candidate.original_filename or "resume",
            "file_type": file_extension,
            "content_type": _get_content_type(file_extension) if file_extension else None,
            "viewer_friendly": file_extension in ['pdf', 'txt', 'doc', 'docx'] if file_extension else False
        }
        
        return response_data
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error viewing resume: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

def _get_content_type(file_extension: str) -> str:
    """Helper function to get content type for file extensions"""
    content_type_map = {
        'pdf': 'application/pdf',
        'doc': 'application/msword',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'txt': 'text/plain',
        'rtf': 'application/rtf'
    }
    return content_type_map.get(file_extension, 'application/octet-stream')

# Shortlisting endpoints
@app.post("/shortlist-by-description/", response_model=ShortlistingResult)
async def shortlist_by_job_description(
    request: Request,
    job_description: str = Form(...),
    min_score: int = Form(70),
    limit: Optional[int] = Form(None),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Shortlist candidates based on a job description (user-specific).
    """
    try:
        if not job_description.strip():
            raise HTTPException(status_code=400, detail="Job description cannot be empty")
        
        result = shortlist_candidates(job_description, min_score, limit, current_user['id'])
        return result
    
    except Exception as e:
        logger.error(f"Error in shortlisting by description: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error shortlisting candidates: {str(e)}")

@app.post("/shortlist-by-file/", response_model=ShortlistingResult)
async def shortlist_by_job_file(
    request: Request,
    file: UploadFile = File(...),
    min_score: int = Form(70),
    limit: Optional[int] = Form(None),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    """
    Shortlist candidates based on a job description file (PDF, DOCX, TXT) (user-specific).
    """
    try:
        # Check if file extension is supported
        file_extension = file.filename.split('.')[-1].lower()
        if file_extension not in ["pdf", "docx", "txt"]:
            raise HTTPException(status_code=400, detail="Unsupported file format. Please upload a PDF, DOCX, or TXT file.")
        
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_file:
            temp_file_path = temp_file.name
            content = await file.read()
            temp_file.write(content)
        
        try:
            # Extract text based on file type
            if file_extension == "pdf":
                job_description = extract_text_from_pdf(temp_file_path)
            elif file_extension == "docx":
                job_description = extract_text_from_docx(temp_file_path)
            elif file_extension == "txt":
                job_description = extract_text_from_txt(temp_file_path)
              # Clean up temporary file
            os.unlink(temp_file_path)
            
            if not job_description.strip():
                raise HTTPException(status_code=400, detail="Could not extract text from the uploaded file")
            
            result = shortlist_candidates(job_description, min_score, limit, current_user['id'])
            return result
        
        except Exception as e:
            # Clean up temporary file in case of error
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
            raise e
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in shortlisting by file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error processing job description file: {str(e)}")

@app.get("/shortlisting-history/")
async def get_shortlisting_history():
    """    Get history of shortlisting operations (placeholder for future implementation).
    """
    return {"message": "Shortlisting history feature coming soon"}

def validate_resume_content(resume_text):
    """
    Use LLM to validate whether the text is actually a resume.
    Returns a dict with validation result and reasoning.
    """
    prompt = f"""You are a smart resume validator. Your job is to determine if the text provided is actually a resume or CV.
    A valid resume should have most of these elements:
    1. Contact information (name, email, phone, etc.)
    2. Education details
    3. Work experience or skills
    4. Some professional information

    Text to validate:
    {resume_text[:4000]}  # Limit text length to avoid token limits

    First, analyze the text and determine if it's a resume or not.
    Then, respond with ONLY a JSON format as follows:
    {{
        "is_resume": true/false,
        "reasoning": "Brief explanation of why this is or isn't a resume",
        "missing_elements": ["List of critical elements missing from the resume, if any"]
    }}

    DO NOT include any other text in your response, ONLY the JSON.
    """

    try:
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": prompt,
                }
            ],
            model="llama3-70b-8192",
            temperature=0.2,
            max_tokens=500
        )
        
        response_text = chat_completion.choices[0].message.content
        
        # Try to parse the JSON response
        import json
        try:
            # Clean the response in case there's any markdown or extra text
            json_text = response_text.strip()
            if json_text.startswith("```json"):
                json_text = json_text.replace("```json", "", 1)
            if json_text.endswith("```"):
                json_text = json_text[:-3]
            
            result = json.loads(json_text)
            return result
        except json.JSONDecodeError:
            logger.error(f"Failed to parse LLM response: {response_text}")
            return {
                "is_resume": False,
                "reasoning": "Error validating resume content",
                "missing_elements": ["Could not analyze resume properly"]
            }
            
    except Exception as e:
        logger.error(f"Error validating resume with LLM: {str(e)}")
        return {
            "is_resume": False,
            "reasoning": f"Error: {str(e)}",
            "missing_elements": ["Validation error occurred"]
        }

@app.post("/validate-resume-content/", response_model=Dict[str, Any])
async def validate_resume_content(
    file: UploadFile = File(...),
    current_user: Dict[str, Any] = Depends(get_current_user_optional)
):
    """
    Validate if a file contains valid resume content using the LLM.
    This endpoint extracts text from the file and uses the LLM to determine if it's a valid resume.
    The LLM also identifies any missing critical elements.
    """
    # Check if file extension is supported
    file_extension = file.filename.split('.')[-1].lower()
    if file_extension not in ["pdf", "docx", "txt"]:
        return {
            "is_resume": False,
            "reasoning": "Unsupported file format. Please upload a PDF, DOCX, or TXT file.",
            "missing_elements": []
        }
    
    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_file:
        temp_file_path = temp_file.name
        content = await file.read()
        temp_file.write(content)
    
    try:
        # Extract text based on file type
        if file_extension == "pdf":
            extracted_text = extract_text_from_pdf(temp_file_path)
        elif file_extension == "docx":
            extracted_text = extract_text_from_docx(temp_file_path)
        elif file_extension == "txt":
            extracted_text = extract_text_from_txt(temp_file_path)
        
        # Check if text extraction succeeded
        if not extracted_text or len(extracted_text.strip()) < 20:  # Minimal text check
            os.unlink(temp_file_path)
            return {
                "is_resume": False,
                "reasoning": "The file appears to be empty or contains too little text to be a valid resume.",
                "missing_elements": ["content"]
            }
            
        # Use LLM to evaluate if this is a valid resume and identify missing elements
        validation_result = validate_resume_with_llm(extracted_text)
        
        # Clean up temporary file
        os.unlink(temp_file_path)
        
        return validation_result
    
    except Exception as e:
        # Clean up temporary file in case of errors
        if os.path.exists(temp_file_path):
            os.unlink(temp_file_path)
        logger.error(f"Error validating resume content: {str(e)}")
        return {
            "is_resume": False,
            "reasoning": f"Error processing file: {str(e)}",
            "missing_elements": []
        }

def validate_resume_with_llm(text: str) -> Dict[str, Any]:
    """
    Use LLM to determine if the text contains a valid resume and identify missing elements.
    
    Args:
        text: The extracted text from the file
        
    Returns:
        dict: Contains validation results with the following keys:
            - is_resume: Boolean indicating if it's a valid resume
            - reasoning: String explaining the decision
            - missing_elements: List of critical elements missing from the resume
    """
    # Create the prompt for LLM
    prompt = f"""Analyze the following text and determine if it is a valid resume. 
A valid resume should contain at least the following critical elements:
1. Contact information (email or phone number)
2. Work experience or education history
3. Skills or qualifications

Text to analyze:
```
{text[:3000]}  # Limit to first 3000 chars for prompt size
```

Respond with a JSON object with the following structure:
{{
  "is_resume": boolean,  # true if it's a valid resume, false otherwise
  "reasoning": string,   # brief explanation for your decision (max 100 words)
  "missing_elements": list of strings  # list of critical elements missing from the resume, empty if none missing
}}

Analyze carefully and be somewhat strict - if the document is clearly not a resume or is missing critical parts that make it unusable for job applications, mark it as not a valid resume."""

    try:
        # Call the LLM with our validation prompt
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": prompt,
                }
            ],
            model="llama3-70b-8192",  # Using LLaMA 3 70B model
            temperature=0.2,  # Lower temperature for consistent responses
            response_format={"type": "json_object"},  # Request JSON response
            max_tokens=500
        )
        
        # Parse the JSON response from the LLM
        response_text = chat_completion.choices[0].message.content
        validation_result = json.loads(response_text)
        
        # Ensure the result has all required fields
        if "is_resume" not in validation_result:
            validation_result["is_resume"] = False
        if "reasoning" not in validation_result:
            validation_result["reasoning"] = "Unable to determine if this is a valid resume."
        if "missing_elements" not in validation_result:
            validation_result["missing_elements"] = []
        
        return validation_result
        
    except Exception as e:
        logger.error(f"Error in LLM validation: {str(e)}")
        # Fallback response in case of LLM error
        return {
            "is_resume": True,  # Default to true in case of error to avoid blocking uploads
            "reasoning": "Automated validation encountered an error, proceeding with basic validation.",
            "missing_elements": []
        }

# Chat Endpoints
@app.post("/chat/start", response_model=Dict[str, str])
async def start_chat_session(
    current_user: Dict[str, Any] = Depends(get_current_user_optional)
):
    """
    Start a new chat session
    """
    try:
        user_id = current_user.get("user_id") if current_user else None
        session_id = chat_service.create_session(user_id)
        
        return {
            "session_id": session_id,
            "message": "Chat session started successfully"
        }
    
    except Exception as e:
        logger.error(f"Error starting chat session: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to start chat session")

@app.post("/chat/message", response_model=ChatResponse)
async def send_chat_message(
    chat_request: ChatRequest,
    current_user: Dict[str, Any] = Depends(get_current_user_optional)
):
    """
    Send a message to the chatbot and get a response
    """
    try:
        user_id = current_user.get("user_id") if current_user else None
        
        # If no session_id provided, create a new session
        if not chat_request.session_id:
            session_id = chat_service.create_session(user_id)
        else:
            session_id = chat_request.session_id
        
        # Generate response using RAG
        response = chat_service.generate_response(
            user_message=chat_request.message,
            session_id=session_id,
            user_id=user_id
        )
        
        return response
    
    except Exception as e:
        logger.error(f"Error processing chat message: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to process chat message")

@app.get("/chat/history/{session_id}", response_model=ChatHistory)
async def get_chat_history(
    session_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user_optional)
):
    """
    Get chat history for a session
    """
    try:
        history = chat_service.get_session_info(session_id)
        
        if not history:
            raise HTTPException(status_code=404, detail="Chat session not found")
        
        return history
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting chat history: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to get chat history")

@app.get("/chat/sessions", response_model=List[Dict[str, Any]])
async def get_user_chat_sessions(
    current_user: Dict[str, Any] = Depends(get_current_user_optional)
):
    """
    Get all chat sessions for the current user
    """
    try:
        user_id = current_user.get("user_id") if current_user else None
        
        # Use the chat service method to get user sessions
        sessions = chat_service.get_user_sessions(user_id)
        return sessions
    
    except Exception as e:
        logger.error(f"Error getting user chat sessions: {str(e)}")
        # Return empty list instead of raising exception for better UX
        return []
